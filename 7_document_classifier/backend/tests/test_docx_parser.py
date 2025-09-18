import pytest
from docx import Document  # Usada para criar o arquivo .docx de teste
import os
from backend.services.docx_parser import DocxParser   

@pytest.fixture
def sample_docx(tmp_path):
    """
    Cria dinamicamente um arquivo .docx de exemplo com texto simples.
    """
    file_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Este é um teste para o parser de docx.")
    doc.add_paragraph("Python é incrível!")
    doc.save(file_path)
    return str(file_path)

class TestDocxParser:
    def test_get_text_from_docx(self, sample_docx):
        """
        Verifica se o texto é extraído corretamente de um arquivo .docx.
        """
        reader = DocxParser(sample_docx)
        text = reader.get_text()

        # Verifica se o texto extraído contém as strings que inserimos.
        assert "Este é um teste para o parser de docx." in text
        assert "Python é incrível!" in text
        # Verifica se o texto é uma string
        assert isinstance(text, str)
        # O método get_text() retorna uma única string, não uma lista.
        assert not isinstance(text, list)

    def test_file_not_found(self, tmp_path):
        """
        Testa se a classe levanta um FileNotFoundError para um arquivo inexistente.
        """
        inexistent_path = tmp_path / "inexistente.docx"
        
        with pytest.raises(FileNotFoundError) as excinfo:
            DocxParser(str(inexistent_path))
        
        assert "O arquivo não foi encontrado" in str(excinfo.value)
        
    def test_empty_docx_file(self, tmp_path):
        """
        Testa o comportamento da classe com um arquivo .docx vazio.
        """
        empty_file_path = tmp_path / "empty.docx"
        Document().save(empty_file_path)
        
        reader = DocxParser(str(empty_file_path))
        text = reader.get_text()
        
        # Um documento .docx vazio deve retornar uma string vazia ou com quebras de linha.
        assert len(text.strip()) == 0